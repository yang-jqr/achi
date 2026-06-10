from docx import Document
from langchain_core.output_parsers import StrOutputParser
from langchain_core.runnables import RunnablePassthrough
from langchain_core.prompts import PromptTemplate

from llm import DeepSeekV4Flash
from prompt import ResumePrompt,ResumePrompt2

# 加载 职位描述
def load_jobs() -> str:
    import os
    job_path = os.path.join(os.path.dirname(__file__), "job.txt")
    with open(job_path, 'r', encoding='utf-8') as f:
        jobs=f.read()
    
    return jobs

def load_doc() -> str:
    """用 python-docx 读取简历（替代废弃的 langchain UnstructuredWordDocumentLoader）"""
    import os
    resume_path = os.path.join(os.path.dirname(__file__), "AI_intern_resume.docx")
    doc = Document(resume_path)
    parts = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    # 也提取表格
    for table in doc.tables:
        for row in table.rows:
            cells = " | ".join(c.text.strip() for c in row.cells)
            if cells.strip():
                parts.append(cells)
    return "\n".join(parts)

def write_resume():
    prompt=PromptTemplate.from_template(ResumePrompt)
    llm=DeepSeekV4Flash()
    chain={
        "input":RunnablePassthrough()
    } | prompt | llm | StrOutputParser()
    ret=chain.invoke(load_jobs())
    print(ret)

def fix_resume():
    prompt=PromptTemplate.from_template(ResumePrompt2)
    llm=DeepSeekV4Flash()
    docs=load_doc()
    chain={
        "resume": lambda _: docs,
        "input":RunnablePassthrough()
    } | prompt | llm | StrOutputParser()
    ret=chain.invoke(load_jobs())
    print(ret)

if __name__ == '__main__':
    fix_resume()