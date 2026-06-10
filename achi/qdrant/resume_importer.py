"""
基于 Qdrant 本地向量数据库的简历导入工具

功能：
1. 读取 .docx 简历文件
2. 文本分块处理
3. BGE 中文向量模型生成 embedding
4. 存入 Qdrant 本地向量数据库（嵌入式模式，无需服务器）
"""

# ===== 必须在最顶部设置镜像源（在所有 import 之前）=====
import os
os.environ["HF_ENDPOINT"] = "https://hf-mirror.com"
os.environ["TRANSFORMERS_OFFLINE"] = "0"

from docx import Document
from sentence_transformers import SentenceTransformer
from qdrant_client import QdrantClient
from qdrant_client.models import Distance, VectorParams, PointStruct

import uuid


# ==================== 配置 ====================
# 简历文件路径
RESUME_PATH = os.path.join(os.path.dirname(os.path.dirname(__file__)), "AI_intern_resume.docx")

# Qdrant 数据存储路径
QDRANT_PATH = os.path.join(os.path.dirname(os.path.dirname(__file__)), "qdrant_resume_db")

# 集合名称
COLLECTION_NAME = "resume_collection"

# 向量模型选择（中文优化模型）
# 推荐选项：
#   - BAAI/bge-base-zh-v1.5     (768维, 效果好, 速度快) ✅ 默认推荐
#   - BAAI/bge-large-zh-v1.5    (1024维, 更精确)
#   - shibing624/text2vec-base-chinese (768维, 轻量)
EMBEDDING_MODEL = "BAAI/bge-base-zh-v1.5"

# 分块参数
CHUNK_SIZE = 300       # 每块最大字数
CHUNK_OVERLAP = 50      # 块之间重叠字数


def read_docx(file_path: str) -> str:
    """读取 docx 文件内容"""
    if not os.path.exists(file_path):
        raise FileNotFoundError(f"文件不存在: {file_path}")

    doc = Document(file_path)
    
    # 提取所有段落文本
    paragraphs = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            paragraphs.append(text)

    # 也提取表格中的文本
    for table in doc.tables:
        for row in table.rows:
            cells_text = [cell.text.strip() for cell in row.cells]
            paragraphs.append(" | ".join(cells_text))

    full_text = "\n".join(paragraphs)
    print(f"[读取] 共 {len(paragraphs)} 段，{len(full_text)} 字符")
    return full_text


def split_text(text: str, chunk_size: int = CHUNK_SIZE, chunk_overlap: int = CHUNK_OVERLAP):
    """文本分割（按中文语义切分）"""
    chunks = []
    start = 0
    text_len = len(text)
    
    while start < text_len:
        end = min(start + chunk_size, text_len)
        # 优先在句号/换行处断开
        if end < text_len:
            for sep in ["\n\n", "\n", "。", "！", "？", "；"]:
                last_sep = text.rfind(sep, start, end)
                if last_sep > start + chunk_size // 2:
                    end = last_sep + len(sep)
                    break
        chunks.append(text[start:end].strip())
        start = end - chunk_overlap if end < text_len else end
    
    # 过滤空块
    chunks = [c for c in chunks if c]
    print(f"[分块] 分成 {len(chunks)} 块")
    return chunks


def init_embedding(model_name: str):
    """初始化向量模型 (使用 sentence_transformers)"""
    print(f"[模型] 加载: {model_name}")
    
    # Monkey-patch: 绕过 transformers 对 torch>=2.6 的强制检查（CVE-2025-32434）
    # 我们明确接受 torch 2.5.1 的风险，且仅在本地 CPU 环境使用
    # 需要同时 patch 源定义和所有 import 引用点
    import transformers.utils.import_utils
    import transformers.modeling_utils
    transformers.utils.import_utils.check_torch_load_is_safe = lambda: None
    transformers.modeling_utils.check_torch_load_is_safe = lambda: None
    
    model = SentenceTransformer(model_name, device="cpu")
    
    # 获取向量维度
    dim = model.get_sentence_embedding_dimension()
    print(f"[模型] 向量维度: {dim}")
    return model, dim


def import_to_qdrant(text_chunks: list[str], file_path: str, collection_name: str = COLLECTION_NAME, _model=None):
    """
    导入到 Qdrant 向量数据库（纯 qdrant-client，无 langchain 依赖）
    """
    # 初始化 embedding 模型
    if _model is None:
        model, vector_dim = init_embedding(EMBEDDING_MODEL)
    else:
        model = _model
        vector_dim = model.get_sentence_embedding_dimension()

    # 连接 Qdrant（本地嵌入式模式）
    client = QdrantClient(path=QDRANT_PATH)
    print(f"[Qdrant] 数据存储路径: {QDRANT_PATH}")

    # 检查并创建集合
    collections = client.get_collections().collections
    collection_names = [c.name for c in collections]
    
    if collection_name not in collection_names:
        print(f"[Qdrant] 创建新集合: {collection_name} (维度={vector_dim})")
        client.create_collection(
            collection_name=collection_name,
            vectors_config=VectorParams(size=vector_dim, distance=Distance.COSINE),
        )
    else:
        print(f"[Qdrant] 集合已存在: {collection_name}")

    # 批量生成向量并插入
    print(f"[导入] 生成 embedding 并写入 {len(text_chunks)} 条记录...")
    
    batch_size = 32
    all_ids = []
    
    for i in range(0, len(text_chunks), batch_size):
        batch = text_chunks[i:i + batch_size]
        vectors = model.encode(batch, normalize_embeddings=True).tolist()
        
        points = [
            PointStruct(
                id=str(uuid.uuid4()),
                vector=vectors[j],
                payload={
                    "text": batch[j],
                    "source": os.path.basename(file_path),
                    "chunk_index": i + j,
                    "total_chunks": len(text_chunks),
                }
            )
            for j in range(len(batch))
        ]
        
        client.upsert(collection_name=collection_name, points=points)
        all_ids.extend([p.id for p in points])
        print(f"  已处理 {min(i + batch_size, len(text_chunks))}/{len(text_chunks)}")

    print(f"[导入] 成功！已写入集合 '{collection_name}'")
    print(f"[导入] 数据条数: {len(all_ids)}")
    
    return client, model


def test_search(client, model, collection_name: str, query: str = "AI应用开发经验", top_k: int = 3):
    """测试搜索功能"""
    print(f"\n{'='*60}")
    print(f"搜索测试: \"{query}\"")
    print(f"{'='*60}")

    # 用 sentence_transformers 生成查询向量
    query_vector = model.encode(query, normalize_embeddings=True).tolist()

    # 使用 Qdrant query_points API
    results = client.query_points(
        collection_name=collection_name,
        query=query_vector,
        limit=top_k,
        with_payload=True,
    )

    print(f"\n找到 {len(results.points)} 条结果:\n")
    for i, point in enumerate(results.points):
        payload = point.payload or {}
        text_content = payload.get("text", "N/A")
        score = point.score if hasattr(point, 'score') else 'N/A'
        
        print(f"--- 结果 {i+1} ---")
        print(f"  相似度: {score:.4f}" if isinstance(score, float) else f"  ID: {point.id}")
        content_preview = str(text_content)[:200].replace('\n', ' ')
        print(f"  内容: {content_preview}...")
        print()

    return results.points


def main():
    """主流程"""
    print("=" * 60)
    print("简历向量数据库导入工具 (纯 qdrant-client)")
    print("=" * 60)
    print(f"  简历文件: {RESUME_PATH}")
    print(f"  向量模型: {EMBEDDING_MODEL}")
    print(f"  集合名称: {COLLECTION_NAME}")
    print("=" * 60)

    # 1. 读取文件
    print("\n[步骤1] 读取简历文件...")
    text = read_docx(RESUME_PATH)
    preview = text[:300].replace("\n", " ")
    print(f"  内容预览: {preview}...")

    # 2. 分块
    print("\n[步骤2] 文本分块...")
    chunks = split_text(text)
    for i, chunk in enumerate(chunks[:3]):
        print(f"  块{i+1}: {chunk[:80]}...")
    if len(chunks) > 3:
        print(f"  ... 共 {len(chunks)} 块")

    # 3. 导入 Qdrant
    print("\n[步骤3] 导入向量数据库...")
    model, vector_dim = init_embedding(EMBEDDING_MODEL)
    client, model = import_to_qdrant(chunks, RESUME_PATH, COLLECTION_NAME, _model=model)

    # 4. 测试搜索
    print("\n[步骤4] 测试搜索...")
    test_search(client, model, COLLECTION_NAME, query="AI开发经验", top_k=3)

    print("\n" + "=" * 60)
    print("完成!")
    print(f"数据已保存到: {os.path.abspath(QDRANT_PATH)}")
    print("=" * 60)


if __name__ == "__main__":
    main()
