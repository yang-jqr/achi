"""
简历文件解析模块
================
支持解析 PDF 和 Word (.docx) 格式的简历文件，提取文本内容。

依赖:
    pip install python-docx PyMuPDF
"""

import os
import re
from typing import Optional


class ResumeParser:
    """
    简历文件解析器

    支持格式:
        - PDF (.pdf) - 使用 PyMuPDF (fitz)
        - Word (.docx) - 使用 python-docx

    用法:
        parser = ResumeParser()
        text = parser.parse("path/to/resume.pdf")
        print(text)
    """

    def parse(self, file_path: str) -> str:
        """
        解析简历文件，提取文本内容

        参数:
            file_path: 简历文件的完整路径

        返回:
            提取出的纯文本内容

        抛出:
            FileNotFoundError: 文件不存在
            ValueError: 不支持的文件格式
        """
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"文件不存在: {file_path}")

        ext = os.path.splitext(file_path)[1].lower()

        if ext == ".pdf":
            return self._parse_pdf(file_path)
        elif ext == ".docx":
            return self._parse_docx(file_path)
        elif ext == ".doc":
            raise ValueError(
                "不支持 .doc 格式（旧版 Word），请先转换为 .docx 格式。\n"
                "转换方法：用 Word 打开文件 -> 另存为 -> 选择 .docx 格式"
            )
        elif ext == ".txt":
            return self._parse_txt(file_path)
        else:
            raise ValueError(
                f"不支持的文件格式: {ext}\n"
                f"支持的格式: .pdf, .docx, .txt"
            )

    def _parse_pdf(self, file_path: str) -> str:
        """解析 PDF 文件"""
        try:
            import fitz  # PyMuPDF
        except ImportError:
            raise ImportError(
                "请先安装 PyMuPDF: uv add PyMuPDF\n"
                "或: pip install PyMuPDF"
            )

        doc = fitz.open(file_path)
        text_parts = []

        for page_num in range(len(doc)):
            page = doc[page_num]
            text = page.get_text()
            if text.strip():
                text_parts.append(text.strip())

        doc.close()
        return self._clean_text("\n\n".join(text_parts))

    def _parse_docx(self, file_path: str) -> str:
        """解析 Word .docx 文件"""
        try:
            from docx import Document
        except ImportError:
            raise ImportError(
                "请先安装 python-docx: uv add python-docx\n"
                "或: pip install python-docx"
            )

        doc = Document(file_path)
        text_parts = []

        for paragraph in doc.paragraphs:
            text = paragraph.text.strip()
            if text:
                text_parts.append(text)

        # 也提取表格内容
        for table in doc.tables:
            for row in table.rows:
                row_texts = []
                for cell in row.cells:
                    cell_text = cell.text.strip()
                    if cell_text:
                        row_texts.append(cell_text)
                if row_texts:
                    text_parts.append(" | ".join(row_texts))

        return self._clean_text("\n".join(text_parts))

    def _parse_txt(self, file_path: str) -> str:
        """解析纯文本文件"""
        with open(file_path, "r", encoding="utf-8") as f:
            return f.read()

    def _clean_text(self, text: str) -> str:
        """清理提取的文本"""
        # 移除多余的空行（保留最多一个连续空行）
        text = re.sub(r"\n{3,}", "\n\n", text)
        # 移除行首行尾的空白
        text = "\n".join(line.strip() for line in text.split("\n"))
        # 移除多余的空格
        text = re.sub(r" {2,}", " ", text)
        return text.strip()


# ============================================================
#  独立运行测试
# ============================================================
if __name__ == "__main__":
    import sys

    if len(sys.argv) < 2:
        print("用法: uv run python resume_parser.py <简历文件路径>")
        print("示例: uv run python resume_parser.py C:/path/to/resume.pdf")
        sys.exit(1)

    file_path = sys.argv[1]
    parser = ResumeParser()

    try:
        text = parser.parse(file_path)
        print(f"✅ 解析成功！共 {len(text)} 个字符")
        print("=" * 60)
        print(text)
        print("=" * 60)
    except Exception as e:
        print(f"❌ 解析失败: {e}")
